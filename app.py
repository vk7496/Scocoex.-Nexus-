import os
import re
import json
import html
import urllib.request
from pathlib import Path
from collections import Counter
from io import BytesIO

import streamlit as st
import pandas as pd
from pypdf import PdfReader
from docx import Document

try:
    from groq import Groq
except ImportError:
    Groq = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
    )
    from reportlab.lib.styles import getSampleStyleSheet
except ImportError:
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    getSampleStyleSheet = None


# ============================================================
# CONFIG
# ============================================================

st.set_page_config(
    page_title="SCOCOEX NEXUS",
    page_icon="🌐",
    layout="wide",
    initial_sidebar_state="expanded",
)

ROOT = Path(__file__).resolve().parent
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = ROOT / "assets"

GITHUB_RAW_BASE = (
    "https://raw.githubusercontent.com/"
    "vk7496/Scocoex.-Nexus-/main/"
)

GITHUB_COMPANIES_URL = GITHUB_RAW_BASE + "companies.json"

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
}


# ============================================================
# LANGUAGES
# ============================================================

LANGUAGES = {
    "English": "en",
    "فارسی": "fa",
    "العربية": "ar",
    "中文": "zh",
    "Русский": "ru",
}


# ============================================================
# TRANSLATIONS
# ============================================================

T = {

    "en": {

        "nav": "NAVIGATION",
        "home": "Home",
        "info": "Information Center",
        "assistant": "AI Assistant",
        "companies": "International Companies",
        "intel": "Intelligence",
        "gallery": "Gallery",
        "about": "About",
        "contact": "Contact",

        "language": "Language",
        "response_language": "AI Response Language",

        "hero_label":
            "GLOBAL INVESTMENT & INTELLIGENCE PLATFORM",

        "hero_title":
            "Connect.<br>Discover.<br>Invest.",

        "hero_text":
            "An AI-powered global intelligence and knowledge "
            "platform designed to connect organizations, "
            "investors, companies and strategic information.",

        "documents": "Documents",
        "chunks": "Knowledge Chunks",
        "international_companies":
            "International Companies",
        "languages": "Languages",

        "knowledge": "Knowledge",

        "knowledge_text":
            "Institutional documents, reports and strategic "
            "information become a searchable knowledge base.",

        "ai": "Artificial Intelligence",

        "ai_text":
            "Groq-powered AI answers questions using information "
            "retrieved from SCOCOEX documents.",

        "network": "Global Network",

        "network_text":
            "Explore international companies, investors and "
            "strategic organizations.",

        "information_title":
            "Information Center",

        "information_text":
            "PDF and DOCX files inside the docs/ folder "
            "are automatically indexed.",

        "search_documents":
            "Search documents",

        "sections":
            "Sections extracted",

        "summary":
            "Generate AI Summary",

        "summary_title":
            "AI Document Summary",

        "ask_title":
            "SCOCOEX AI Assistant",

        "ask_caption":
            "Ask questions about the information contained "
            "in the SCOCOEX Knowledge Base.",

        "your_question":
            "Your question",

        "ask_button":
            "Ask SCOCOEX AI",

        "searching":
            "Searching the SCOCOEX Knowledge Base...",

        "generating":
            "Generating intelligence...",

        "sources":
            "Sources",

        "companies_title":
            "International Companies",

        "search_company":
            "Search company",

        "visit":
            "Visit official website →",

        "company_count":
            "organizations",

        "intel_title":
            "Intelligence Dashboard",

        "intel_caption":
            "Document intelligence and numerical overview.",

        "document_size":
            "Document Size",

        "numeric":
            "Numeric Information",

        "about_title":
            "About SCOCOEX",

        "contact_title":
            "Connect with SCOCOEX",

        "name":
            "Name",

        "company_org":
            "Company / Organization",

        "email":
            "Email",

        "interest":
            "Interest",

        "message":
            "Message",

        "send":
            "Send Inquiry",

        "received":
            "Thank you. Your inquiry has been received.",

        "gallery_title":
            "SCOCOEX Gallery",

        "gallery_text":
            "Event images can be added here later.",

        "no_docs":
            "No PDF or DOCX files were found inside docs.",

        "docs_missing":
            "The docs folder does not exist. "
            "Create a folder named docs in the repository.",

        "no_companies":
            "No companies were loaded.",

        "companies_error":
            "companies.json could not be loaded.",

        "groq_missing":
            "Groq API is not configured. "
            "Add GROQ_API_KEY in Streamlit Secrets.",

        "no_results":
            "I could not find relevant information "
            "in the current Knowledge Base.",

        "no_text":
            "No readable text was found.",

        "download_pdf":
            "Download AI Answer as PDF",

        "source_file":
            "Source",

        "words":
            "Words",

        "numeric_values":
            "Numeric values",

        "extracted_sections":
            "Extracted sections",

        "sector":
            "Sector",

        "company_description":
            "Description",

        "all_companies":
            "All organizations",

        "status":
            "Status",

        "loaded":
            "Loaded",

        "local":
            "Local",

        "github":
            "GitHub",

        "country":
            "Country",

        "all":
            "All",

        "strategic":
            "Strategic Organizations",

        "international_network":
            "International Company Network",

        "search_network":
            "Search organizations, sectors or countries...",

        "organizations_in_network":
            "organizations in network",

    },


    "fa": {

        "nav": "منوی اصلی",
        "home": "خانه",
        "info": "مرکز اطلاعات",
        "assistant": "دستیار هوش مصنوعی",
        "companies": "شرکت‌های بین‌المللی",
        "intel": "داشبورد هوشمندی",
        "gallery": "گالری",
        "about": "درباره ما",
        "contact": "ارتباط با ما",

        "language": "زبان",
        "response_language":
            "زبان پاسخ هوش مصنوعی",

        "hero_label":
            "پلتفرم جهانی سرمایه‌گذاری و هوشمندی",

        "hero_title":
            "ارتباط.<br>کشف.<br>سرمایه‌گذاری.",

        "hero_text":
            "یک پلتفرم جهانی مبتنی بر هوش مصنوعی برای "
            "اتصال سازمان‌ها، سرمایه‌گذاران، شرکت‌ها و "
            "اطلاعات راهبردی.",

        "documents": "اسناد",
        "chunks": "قطعات دانش",
        "international_companies":
            "شرکت‌های بین‌المللی",
        "languages": "زبان‌ها",

        "knowledge": "دانش",

        "knowledge_text":
            "اسناد سازمانی، گزارش‌ها و اطلاعات راهبردی "
            "به یک پایگاه دانش قابل جستجو تبدیل می‌شوند.",

        "ai": "هوش مصنوعی",

        "ai_text":
            "هوش مصنوعی مبتنی بر Groq با استفاده از "
            "اطلاعات بازیابی‌شده از اسناد SCOCOEX پاسخ می‌دهد.",

        "network": "شبکه جهانی",

        "network_text":
            "شرکت‌ها، سرمایه‌گذاران و سازمان‌های "
            "راهبردی بین‌المللی را بررسی کنید.",

        "information_title":
            "مرکز اطلاعات",

        "information_text":
            "فایل‌های PDF و DOCX موجود در پوشه docs "
            "به‌صورت خودکار خوانده و ایندکس می‌شوند.",

        "search_documents":
            "جستجو در اسناد",

        "sections":
            "بخش‌های استخراج‌شده",

        "summary":
            "ساخت خلاصه هوش مصنوعی",

        "summary_title":
            "خلاصه هوشمند سند",

        "ask_title":
            "دستیار هوشمند SCOCOEX",

        "ask_caption":
            "درباره اطلاعات موجود در پایگاه دانش "
            "SCOCOEX سؤال بپرسید.",

        "your_question":
            "سؤال شما",

        "ask_button":
            "پرسش از SCOCOEX AI",

        "searching":
            "در حال جستجو در پایگاه دانش SCOCOEX...",

        "generating":
            "در حال تولید تحلیل...",

        "sources":
            "منابع",

        "companies_title":
            "شرکت‌های بین‌المللی",

        "search_company":
            "جستجوی شرکت",

        "visit":
            "مشاهده وب‌سایت رسمی ←",

        "company_count":
            "سازمان",

        "intel_title":
            "داشبورد هوشمندی",

        "intel_caption":
            "تحلیل اسناد و مرور اطلاعات عددی.",

        "document_size":
            "حجم اسناد",

        "numeric":
            "اطلاعات عددی",

        "about_title":
            "درباره SCOCOEX",

        "contact_title":
            "ارتباط با SCOCOEX",

        "name":
            "نام",

        "company_org":
            "شرکت / سازمان",

        "email":
            "ایمیل",

        "interest":
            "موضوع همکاری",

        "message":
            "پیام",

        "send":
            "ارسال درخواست",

        "received":
            "درخواست شما با موفقیت دریافت شد.",

        "gallery_title":
            "گالری SCOCOEX",

        "gallery_text":
            "تصاویر رویداد را می‌توان در این بخش اضافه کرد.",

        "no_docs":
            "هیچ فایل PDF یا DOCX در پوشه docs پیدا نشد.",

        "docs_missing":
            "پوشه docs وجود ندارد. یک پوشه با نام docs "
            "در Repository بسازید.",

        "no_companies":
            "هیچ شرکتی بارگذاری نشد.",

        "companies_error":
            "فایل companies.json قابل بارگذاری نیست.",

        "groq_missing":
            "کلید Groq تنظیم نشده است. "
            "GROQ_API_KEY را در Streamlit Secrets قرار دهید.",

        "no_results":
            "اطلاعات مرتبطی در پایگاه دانش فعلی پیدا نشد.",

        "no_text":
            "متن قابل خواندن پیدا نشد.",

        "download_pdf":
            "دریافت پاسخ AI به صورت PDF",

        "source_file":
            "منبع",

        "words":
            "کلمات",

        "numeric_values":
            "مقادیر عددی",

        "extracted_sections":
            "بخش‌های استخراج‌شده",

        "sector":
            "حوزه",

        "company_description":
            "توضیحات",

        "all_companies":
            "همه سازمان‌ها",

        "status":
            "وضعیت",

        "loaded":
            "بارگذاری شد",

        "local":
            "محلی",

        "github":
            "گیت‌هاب",

        "country":
            "کشور",

        "all":
            "همه",

        "strategic":
            "سازمان‌های راهبردی",

        "international_network":
            "شبکه شرکت‌های بین‌المللی",

        "search_network":
            "جستجوی سازمان، حوزه یا کشور...",

        "organizations_in_network":
            "سازمان در شبکه",

    },


    "ar": {

        "nav": "التنقل",
        "home": "الرئيسية",
        "info": "مركز المعلومات",
        "assistant": "المساعد الذكي",
        "companies": "الشركات الدولية",
        "intel": "لوحة الذكاء",
        "gallery": "المعرض",
        "about": "من نحن",
        "contact": "تواصل معنا",

        "language": "اللغة",
        "response_language":
            "لغة رد الذكاء الاصطناعي",

        "hero_label":
            "منصة عالمية للاستثمار والذكاء",

        "hero_title":
            "تواصل.<br>اكتشف.<br>استثمر.",

        "hero_text":
            "منصة عالمية مدعومة بالذكاء الاصطناعي "
            "لربط المؤسسات والمستثمرين والشركات "
            "والمعلومات الاستراتيجية.",

        "documents": "الوثائق",
        "chunks": "مقاطع المعرفة",
        "international_companies":
            "الشركات الدولية",
        "languages": "اللغات",

        "knowledge": "المعرفة",

        "knowledge_text":
            "تحويل الوثائق والتقارير والمعلومات "
            "الاستراتيجية إلى قاعدة معرفة قابلة للبحث.",

        "ai": "الذكاء الاصطناعي",

        "ai_text":
            "يجيب الذكاء الاصطناعي المدعوم من Groq "
            "باستخدام المعلومات المسترجعة من وثائق SCOCOEX.",

        "network": "الشبكة العالمية",

        "network_text":
            "استكشف الشركات والمستثمرين والمنظمات "
            "الاستراتيجية الدولية.",

        "information_title":
            "مركز المعلومات",

        "information_text":
            "يتم فهرسة ملفات PDF وDOCX الموجودة "
            "في مجلد docs تلقائياً.",

        "search_documents":
            "البحث في الوثائق",

        "sections":
            "الأقسام المستخرجة",

        "summary":
            "إنشاء ملخص بالذكاء الاصطناعي",

        "summary_title":
            "ملخص ذكي للوثيقة",

        "ask_title":
            "مساعد SCOCOEX الذكي",

        "ask_caption":
            "اطرح أسئلة حول المعلومات الموجودة "
            "في قاعدة معرفة SCOCOEX.",

        "your_question":
            "سؤالك",

        "ask_button":
            "اسأل SCOCOEX AI",

        "searching":
            "جارٍ البحث في قاعدة المعرفة...",

        "generating":
            "جارٍ إنشاء التحليل...",

        "sources":
            "المصادر",

        "companies_title":
            "الشركات الدولية",

        "search_company":
            "البحث عن شركة",

        "visit":
            "زيارة الموقع الرسمي ←",

        "company_count":
            "منظمة",

        "intel_title":
            "لوحة الذكاء",

        "intel_caption":
            "تحليل الوثائق ومراجعة المعلومات الرقمية.",

        "document_size":
            "حجم الوثائق",

        "numeric":
            "المعلومات الرقمية",

        "about_title":
            "عن SCOCOEX",

        "contact_title":
            "تواصل مع SCOCOEX",

        "name":
            "الاسم",

        "company_org":
            "الشركة / المؤسسة",

        "email":
            "البريد الإلكتروني",

        "interest":
            "مجال الاهتمام",

        "message":
            "الرسالة",

        "send":
            "إرسال الطلب",

        "received":
            "تم استلام طلبك بنجاح.",

        "gallery_title":
            "معرض SCOCOEX",

        "gallery_text":
            "يمكن إضافة صور الفعالية هنا.",

        "no_docs":
            "لم يتم العثور على ملفات PDF أو DOCX داخل docs.",

        "docs_missing":
            "مجلد docs غير موجود.",

        "no_companies":
            "لم يتم تحميل شركات.",

        "companies_error":
            "تعذر تحميل companies.json.",

        "groq_missing":
            "لم يتم إعداد Groq API. أضف GROQ_API_KEY "
            "في Streamlit Secrets.",

        "no_results":
            "لم أجد معلومات ذات صلة في قاعدة المعرفة الحالية.",

        "no_text":
            "لم يتم العثور على نص قابل للقراءة.",

        "download_pdf":
            "تنزيل إجابة AI بصيغة PDF",

        "source_file":
            "المصدر",

        "words":
            "الكلمات",

        "numeric_values":
            "القيم الرقمية",

        "extracted_sections":
            "الأقسام المستخرجة",

        "sector":
            "القطاع",

        "company_description":
            "الوصف",

        "all_companies":
            "جميع المؤسسات",

        "status":
            "الحالة",

        "loaded":
            "تم التحميل",

        "local":
            "محلي",

        "github":
            "GitHub",

        "country":
            "الدولة",

        "all":
            "الكل",

        "strategic":
            "المنظمات الاستراتيجية",

        "international_network":
            "شبكة الشركات الدولية",

        "search_network":
            "البحث عن المؤسسات أو القطاعات أو الدول...",

        "organizations_in_network":
            "منظمة في الشبكة",

    },


    "zh": {

        "nav": "导航",
        "home": "首页",
        "info": "信息中心",
        "assistant": "AI 助手",
        "companies": "国际企业",
        "intel": "智能分析",
        "gallery": "图库",
        "about": "关于我们",
        "contact": "联系我们",

        "language": "语言",
        "response_language":
            "AI 回复语言",

        "hero_label":
            "全球投资与智能平台",

        "hero_title":
            "连接。<br>探索。<br>投资。",

        "hero_text":
            "一个由人工智能驱动的全球知识与智能平台，"
            "用于连接组织、投资者、企业和战略信息。",

        "documents": "文件",
        "chunks": "知识片段",
        "international_companies":
            "国际企业",
        "languages": "语言",

        "knowledge": "知识",

        "knowledge_text":
            "将机构文件、报告和战略信息转化为可搜索的知识库。",

        "ai": "人工智能",

        "ai_text":
            "由 Groq 驱动的 AI 使用 SCOCOEX 文件中的检索信息回答问题。",

        "network": "全球网络",

        "network_text":
            "探索国际企业、投资者和战略组织。",

        "information_title":
            "信息中心",

        "information_text":
            "docs/ 文件夹中的 PDF 和 DOCX 会自动建立索引。",

        "search_documents":
            "搜索文件",

        "sections":
            "提取章节",

        "summary":
            "生成 AI 摘要",

        "summary_title":
            "AI 文件摘要",

        "ask_title":
            "SCOCOEX AI 助手",

        "ask_caption":
            "询问 SCOCOEX 知识库中的信息。",

        "your_question":
            "您的问题",

        "ask_button":
            "询问 SCOCOEX AI",

        "searching":
            "正在搜索知识库...",

        "generating":
            "正在生成分析...",

        "sources":
            "来源",

        "companies_title":
            "国际企业",

        "search_company":
            "搜索企业",

        "visit":
            "访问官方网站 →",

        "company_count":
            "家机构",

        "intel_title":
            "智能分析仪表板",

        "intel_caption":
            "文件智能分析和数字信息概览。",

        "document_size":
            "文件规模",

        "numeric":
            "数字信息",

        "about_title":
            "关于 SCOCOEX",

        "contact_title":
            "联系我们",

        "name":
            "姓名",

        "company_org":
            "公司 / 组织",

        "email":
            "邮箱",

        "interest":
            "合作方向",

        "message":
            "留言",

        "send":
            "发送请求",

        "received":
            "您的请求已收到。",

        "gallery_title":
            "SCOCOEX 图库",

        "gallery_text":
            "之后可以在这里加入活动图片。",

        "no_docs":
            "docs/ 中没有找到 PDF 或 DOCX 文件。",

        "docs_missing":
            "docs 文件夹不存在。",

        "no_companies":
            "未加载企业。",

        "companies_error":
            "无法加载 companies.json。",

        "groq_missing":
            "尚未配置 Groq API。请在 Streamlit Secrets 中添加 GROQ_API_KEY。",

        "no_results":
            "当前知识库中没有找到相关信息。",

        "no_text":
            "没有找到可读取的文本。",

        "download_pdf":
            "下载 AI 回答 PDF",

        "source_file":
            "来源",

        "words":
            "词数",

        "numeric_values":
            "数字值",

        "extracted_sections":
            "提取章节",

        "sector":
            "领域",

        "company_description":
            "描述",

        "all_companies":
            "所有机构",

        "status":
            "状态",

        "loaded":
            "已加载",

        "local":
            "本地",

        "github":
            "GitHub",

        "country":
            "国家",

        "all":
            "全部",

        "strategic":
            "战略组织",

        "international_network":
            "国际企业网络",

        "search_network":
            "搜索组织、领域或国家...",

        "organizations_in_network":
            "网络中的组织",

    },


    "ru": {

        "nav": "НАВИГАЦИЯ",
        "home": "Главная",
        "info": "Информационный центр",
        "assistant": "AI-ассистент",
        "companies": "Международные компании",
        "intel": "Интеллект",
        "gallery": "Галерея",
        "about": "О нас",
        "contact": "Контакты",

        "language": "Язык",
        "response_language":
            "Язык ответа AI",

        "hero_label":
            "ГЛОБАЛЬНАЯ ПЛАТФОРМА ИНВЕСТИЦИЙ И АНАЛИТИКИ",

        "hero_title":
            "Связывай.<br>Исследуй.<br>Инвестируй.",

        "hero_text":
            "Глобальная интеллектуальная платформа на базе ИИ "
            "для объединения организаций, инвесторов, компаний "
            "и стратегической информации.",

        "documents": "Документы",
        "chunks": "Фрагменты знаний",
        "international_companies":
            "Международные компании",
        "languages": "Языки",

        "knowledge": "Знания",

        "knowledge_text":
            "Документы, отчёты и стратегическая информация "
            "превращаются в поисковую базу знаний.",

        "ai":
            "Искусственный интеллект",

        "ai_text":
            "AI на базе Groq отвечает на вопросы, используя "
            "информацию из документов SCOCOEX.",

        "network":
            "Глобальная сеть",

        "network_text":
            "Исследуйте международные компании, инвесторов "
            "и стратегические организации.",

        "information_title":
            "Информационный центр",

        "information_text":
            "PDF и DOCX в папке docs автоматически индексируются.",

        "search_documents":
            "Поиск документов",

        "sections":
            "Извлечённые разделы",

        "summary":
            "Создать AI-резюме",

        "summary_title":
            "AI-резюме документа",

        "ask_title":
            "AI-ассистент SCOCOEX",

        "ask_caption":
            "Задавайте вопросы по информации из базы знаний SCOCOEX.",

        "your_question":
            "Ваш вопрос",

        "ask_button":
            "Спросить SCOCOEX AI",

        "searching":
            "Поиск в базе знаний...",

        "generating":
            "Создание анализа...",

        "sources":
            "Источники",

        "companies_title":
            "Международные компании",

        "search_company":
            "Поиск компании",

        "visit":
            "Официальный сайт →",

        "company_count":
            "организаций",

        "intel_title":
            "Интеллектуальная панель",

        "intel_caption":
            "Анализ документов и обзор числовой информации.",

        "document_size":
            "Размер документов",

        "numeric":
            "Числовая информация",

        "about_title":
            "О SCOCOEX",

        "contact_title":
            "Связаться с SCOCOEX",

        "name":
            "Имя",

        "company_org":
            "Компания / организация",

        "email":
            "Email",

        "interest":
            "Направление интереса",

        "message":
            "Сообщение",

        "send":
            "Отправить запрос",

        "received":
            "Ваш запрос получен.",

        "gallery_title":
            "Галерея SCOCOEX",

        "gallery_text":
            "Позже здесь можно добавить изображения мероприятия.",

        "no_docs":
            "В папке docs не найдено PDF или DOCX.",

        "docs_missing":
            "Папка docs не существует.",

        "no_companies":
            "Компании не загружены.",

        "companies_error":
            "Не удалось загрузить companies.json.",

        "groq_missing":
            "Groq API ещё не настроен. Добавьте GROQ_API_KEY в Streamlit Secrets.",

        "no_results":
            "В текущей базе знаний не найдена релевантная информация.",

        "no_text":
            "Не найден читаемый текст.",

        "download_pdf":
            "Скачать ответ AI в PDF",

        "source_file":
            "Источник",

        "words":
            "Слова",

        "numeric_values":
            "Числовые значения",

        "extracted_sections":
            "Извлечённые разделы",

        "sector":
            "Сектор",

        "company_description":
            "Описание",

        "all_companies":
            "Все организации",

        "status":
            "Статус",

        "loaded":
            "Загружено",

        "local":
            "Локально",

        "github":
            "GitHub",

        "country":
            "Страна",

        "all":
            "Все",

        "strategic":
            "Стратегические организации",

        "international_network":
            "Международная сеть компаний",

        "search_network":
            "Поиск организаций, отраслей или стран...",

        "organizations_in_network":
            "организаций в сети",

    },
}


# ============================================================
# SESSION
# ============================================================

if "language_name" not in st.session_state:
    st.session_state.language_name = "English"

if "lang_code" not in st.session_state:
    st.session_state.lang_code = "en"


def t(key):
    lang = st.session_state.get("lang_code", "en")
    return T.get(lang, T["en"]).get(key, key)


def esc(value):
    return html.escape(str(value or ""))


def direction():
    return (
        "rtl"
        if st.session_state.get("lang_code") in {"fa", "ar"}
        else "ltr"
    )


# ============================================================
# CSS
# ============================================================

st.markdown(
    """
<style>

@import url(
'https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap'
);

html,
body,
[class*="css"] {
    font-family: Inter, sans-serif;
}

[data-testid="stAppViewContainer"] {
    background:
        radial-gradient(
            circle at 8% 0%,
            rgba(84,76,255,.18),
            transparent 28%
        ),
        radial-gradient(
            circle at 95% 8%,
            rgba(0,190,255,.10),
            transparent 25%
        ),
        #070A12;
}

[data-testid="stHeader"] {
    background: transparent !important;
}

.block-container {
    max-width: 1450px;
    padding-top: 2rem;
    padding-bottom: 4rem;
}


/* ============================================================
   SIDEBAR
   ============================================================ */

[data-testid="stSidebar"] {
    background:
        radial-gradient(
            circle at 20% 0%,
            rgba(102,92,255,.20),
            transparent 32%
        ),
        linear-gradient(
            180deg,
            #080B16 0%,
            #0D1222 55%,
            #080B14 100%
        ) !important;

    border-right:
        1px solid rgba(255,255,255,.09) !important;
}

[data-testid="stSidebar"] * {
    color: #F4F6FC !important;
}

[data-testid="stSidebar"] label {
    color: #D5DBEA !important;
    font-weight: 600 !important;
}

[data-testid="stSidebar"] .stCaption {
    color: #8E99B2 !important;
}

[data-testid="stSidebar"] [data-baseweb="select"] {
    background:
        rgba(255,255,255,.065) !important;

    border:
        1px solid rgba(255,255,255,.13) !important;

    border-radius:
        12px !important;
}

[data-testid="stSidebar"]
[data-baseweb="select"] * {
    color: #FFFFFF !important;
}

[data-testid="stSidebar"]
[data-testid="stRadio"] label {
    padding:
        8px 10px !important;

    border-radius:
        11px !important;

    transition:
        all .18s ease;
}

[data-testid="stSidebar"]
[data-testid="stRadio"] label:hover {
    background:
        rgba(145,132,255,.14) !important;
}


/* ============================================================
   HERO
   ============================================================ */

.hero {
    padding:
        55px 45px;

    border-radius:
        30px;

    border:
        1px solid rgba(255,255,255,.10);

    background:
        radial-gradient(
            circle at 80% 20%,
            rgba(105,91,255,.20),
            transparent 32%
        ),
        linear-gradient(
            135deg,
            rgba(27,31,57,.98),
            rgba(9,12,23,.96)
        );

    box-shadow:
        0 30px 90px rgba(0,0,0,.30);

    margin-bottom:
        30px;
}

.hero-label {
    color:
        #A6B0FF;

    font-size:
        12px;

    font-weight:
        800;

    letter-spacing:
        .18em;

    text-transform:
        uppercase;
}

.hero h1 {
    font-size:
        clamp(44px, 7vw, 84px);

    line-height:
        .98;

    margin:
        16px 0 20px;

    font-weight:
        800;
}

.hero p {
    color:
        #B8C0D4;

    font-size:
        17px;

    max-width:
        850px;

    line-height:
        1.8;
}


/* ============================================================
   CARDS
   ============================================================ */

.metric-card,
.feature-card {
    padding:
        22px;

    border-radius:
        22px;

    background:
        rgba(255,255,255,.045);

    border:
        1px solid rgba(255,255,255,.08);

    height:
        100%;
}

.metric-number {
    font-size:
        30px;

    font-weight:
        800;

    margin-top:
        8px;
}

.muted {
    color:
        #8F99AE;

    font-size:
        13px;
}

.answer-box {
    padding:
        27px;

    border-radius:
        22px;

    background:
        rgba(255,255,255,.045);

    border:
        1px solid rgba(255,255,255,.08);

    line-height:
        1.85;

    box-shadow:
        0 18px 45px rgba(0,0,0,.12);
}

.section-title {
    font-size:
        30px;

    font-weight:
        800;

    margin-bottom:
        8px;
}


/* ============================================================
   STRATEGIC ORGANIZATIONS
   ============================================================ */

.strategic-header,
.directory-header {
    font-size:
        12px;

    font-weight:
        800;

    letter-spacing:
        .14em;

    color:
        #AEB7FF;

    margin:
        28px 0 12px;
}

.strategic-header span,
.directory-header span {
    color:
        #C9A86A;

    margin-right:
        7px;
}

.strategic-card {
    min-height:
        245px;

    padding:
        24px;

    border-radius:
        24px;

    border:
        1px solid rgba(255,255,255,.10);

    background:
        linear-gradient(
            145deg,
            rgba(255,255,255,.075),
            rgba(255,255,255,.025)
        );

    box-shadow:
        0 20px 55px rgba(0,0,0,.22);

    transition:
        transform .2s ease,
        border-color .2s ease;
}

.strategic-card:hover {
    transform:
        translateY(-5px);

    border-color:
        rgba(255,255,255,.22);
}

.strategic-gold {
    box-shadow:
        0 20px 55px rgba(205,169,91,.12);
}

.strategic-green {
    box-shadow:
        0 20px 55px rgba(71,184,132,.10);
}

.strategic-blue {
    box-shadow:
        0 20px 55px rgba(75,137,255,.12);
}

.strategic-top {
    display:
        flex;

    gap:
        14px;

    align-items:
        center;
}

.strategic-logo,
.strategic-logo-fallback {
    width:
        62px;

    height:
        62px;

    border-radius:
        17px;

    object-fit:
        contain;

    background:
        #FFFFFF;

    padding:
        7px;

    flex:
        0 0 auto;
}

.strategic-logo-fallback {
    display:
        flex;

    align-items:
        center;

    justify-content:
        center;

    background:
        rgba(255,255,255,.10);

    color:
        #FFFFFF;

    font-size:
        24px;

    font-weight:
        800;
}

.strategic-kicker {
    font-size:
        9px;

    letter-spacing:
        .16em;

    color:
        #8E99B2;

    font-weight:
        800;
}

.strategic-name {
    font-size:
        17px;

    font-weight:
        800;

    margin-top:
        5px;

    line-height:
        1.25;
}

.strategic-sector {
    margin-top:
        22px;

    color:
        #D7DCF0;

    font-size:
        12px;

    font-weight:
        600;
}

.strategic-copy {
    margin-top:
        13px;

    color:
        #8F9AB2;

    font-size:
        12px;

    line-height:
        1.65;
}

.strategic-link {
    display:
        inline-block;

    margin-top:
        18px;

    color:
        #B9C0FF !important;

    text-decoration:
        none !important;

    font-size:
        12px;

    font-weight:
        700;
}


/* ============================================================
   COMPANY DIRECTORY
   ============================================================ */

.directory-count {
    margin:
        14px 0 16px;

    font-size:
        18px;

    font-weight:
        800;
}

.directory-count span {
    color:
        #7F8AA2;

    font-size:
        12px;

    font-weight:
        500;
}

.directory-card {
    min-height:
        175px;

    margin-bottom:
        16px;

    padding:
        19px;

    border-radius:
        20px;

    background:
        rgba(255,255,255,.045);

    border:
        1px solid rgba(255,255,255,.075);

    transition:
        transform .18s ease,
        border-color .18s ease,
        background .18s ease;
}

.directory-card:hover {
    transform:
        translateY(-3px);

    background:
        rgba(255,255,255,.065);

    border-color:
        rgba(166,176,255,.22);
}

.directory-top {
    display:
        flex;

    align-items:
        center;

    gap:
        13px;
}

.directory-logo,
.directory-initial {
    width:
        50px;

    height:
        50px;

    border-radius:
        14px;

    flex:
        0 0 auto;

    object-fit:
        contain;

    background:
        #FFFFFF;

    padding:
        6px;
}

.directory-initial {
    display:
        flex;

    align-items:
        center;

    justify-content:
        center;

    background:
        linear-gradient(
            135deg,
            #202744,
            #12172A
        );

    color:
        #DDE2FF;

    font-size:
        20px;

    font-weight:
        800;
}

.directory-name {
    font-size:
        15px;

    font-weight:
        750;

    line-height:
        1.35;
}

.directory-meta {
    color:
        #8792AA;

    font-size:
        11px;

    margin-top:
        14px;

    min-height:
        30px;

    line-height:
        1.5;
}

.directory-bottom {
    margin-top:
        12px;

    border-top:
        1px solid rgba(255,255,255,.06);

    padding-top:
        10px;
}

.directory-link {
    color:
        #AEB6FF !important;

    text-decoration:
        none !important;

    font-size:
        11px;

    font-weight:
        700;
}


/* ============================================================
   RTL
   ============================================================ */

[dir="rtl"] .hero,
[dir="rtl"] .answer-box,
[dir="rtl"] .feature-card,
[dir="rtl"] .metric-card,
[dir="rtl"] .strategic-card,
[dir="rtl"] .directory-card {
    text-align:
        right;
}


/* ============================================================
   MOBILE
   ============================================================ */

@media (max-width: 700px) {

    .hero {
        padding:
            36px 24px;
    }

    .hero h1 {
        font-size:
            46px;
    }

}

</style>
""",
    unsafe_allow_html=True,
)


# ============================================================
# LANGUAGE SELECTOR
# ============================================================

selected_language = st.sidebar.selectbox(
    t("language"),
    list(LANGUAGES.keys()),
    index=list(LANGUAGES.keys()).index(
        st.session_state.language_name
    ),
    key="sidebar_language_selector_v4",
)

if selected_language != st.session_state.language_name:

    st.session_state.language_name = (
        selected_language
    )

    st.session_state.lang_code = (
        LANGUAGES[selected_language]
    )

    st.rerun()


st.sidebar.markdown(
    "### 🌐 SCOCOEX NEXUS"
)


# ============================================================
# DOCUMENT EXTRACTION
# ============================================================

def extract_pdf(path):

    sections = []

    reader = PdfReader(str(path))

    for page_no, page in enumerate(
        reader.pages,
        start=1
    ):

        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""

        if text.strip():

            sections.append(
                {
                    "page": page_no,
                    "text": text.strip(),
                }
            )

    return sections


def extract_docx(path):

    document = Document(str(path))

    sections = []

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:

            sections.append(
                {
                    "page": None,
                    "text": text,
                }
            )

    for table_no, table in enumerate(
        document.tables,
        start=1
    ):

        rows = []

        for row in table.rows:

            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            rows.append(
                " | ".join(cells)
            )

        if rows:

            sections.append(
                {
                    "page": None,
                    "text":
                        f"[TABLE {table_no}]\n"
                        + "\n".join(rows),
                }
            )

    return sections


@st.cache_data(show_spinner=False)
def load_all_documents():

    documents = []

    if not DOCS_DIR.exists():
        return documents

    for path in sorted(
        DOCS_DIR.rglob("*")
    ):

        if not path.is_file():
            continue

        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        try:

            if path.suffix.lower() == ".pdf":
                sections = extract_pdf(path)
            else:
                sections = extract_docx(path)

            documents.append(
                {
                    "name": path.name,
                    "path": str(path),
                    "extension":
                        path.suffix.lower(),
                    "sections": sections,
                    "error": None,
                }
            )

        except Exception as exc:

            documents.append(
                {
                    "name": path.name,
                    "path": str(path),
                    "extension":
                        path.suffix.lower(),
                    "sections": [],
                    "error": str(exc),
                }
            )

    return documents


def normalize_text(text):

    return re.sub(
        r"\s+",
        " ",
        text or ""
    ).strip()


@st.cache_data(show_spinner=False)
def create_chunks(
    documents,
    chunk_size=1400,
    overlap=200
):

    chunks = []

    step = max(
        1,
        chunk_size - overlap
    )

    for document in documents:

        for section in document.get(
            "sections",
            []
        ):

            text = normalize_text(
                section.get(
                    "text",
                    ""
                )
            )

            if not text:
                continue

            start = 0

            while start < len(text):

                chunks.append(
                    {
                        "document":
                            document["name"],

                        "page":
                            section.get("page"),

                        "text":
                            text[
                                start:
                                start + chunk_size
                            ],
                    }
                )

                start += step

    return chunks


# ============================================================
# SEARCH / RETRIEVAL
# ============================================================

def tokenize(text):

    return re.findall(
        r"[\w\u0600-\u06FF\u4e00-\u9fff\u0400-\u04FF]+",
        (text or "").lower(),
    )


def retrieve_chunks(
    question,
    chunks,
    top_k=7
):

    q_tokens = set(
        tokenize(question)
    )

    if not q_tokens:
        return []

    scored = []

    for chunk in chunks:

        tokens = tokenize(
            chunk["text"]
        )

        counts = Counter(tokens)

        overlap = sum(
            counts.get(tok, 0)
            for tok in q_tokens
        )

        unique_overlap = len(
            q_tokens.intersection(
                counts.keys()
            )
        )

        if overlap:

            score = (
                overlap
                + unique_overlap * 2
            )

            scored.append(
                (
                    score,
                    chunk
                )
            )

    scored.sort(
        key=lambda item: item[0],
        reverse=True
    )

    return [
        chunk
        for _, chunk in scored[:top_k]
    ]


# ============================================================
# COMPANY DATA
# ============================================================

def parse_companies_payload(payload):

    if isinstance(payload, dict):

        for key in (
            "companies",
            "organizations",
            "data",
            "items"
        ):

            if isinstance(
                payload.get(key),
                list
            ):

                payload = payload[key]
                break

    if not isinstance(
        payload,
        list
    ):

        raise ValueError(
            "companies.json must contain "
            "a JSON array."
        )

    cleaned = []

    for item in payload:

        if not isinstance(
            item,
            dict
        ):
            continue

        company = dict(item)

        if (
            not company.get("logo")
            and company.get("logo_url")
        ):

            company["logo"] = (
                company["logo_url"]
            )

        cleaned.append(company)

    return cleaned


@st.cache_data(show_spinner=False)
def load_companies():

    local_candidates = [

        ROOT / "companies.json",

        ROOT / "companies. json",

        ROOT / "knowledge_base" /
        "companies.json",

    ]

    local_error = (
        "Local companies.json not found."
    )

    for path in local_candidates:

        if path.exists():

            try:

                data = json.loads(
                    path.read_text(
                        encoding="utf-8"
                    )
                )

                return (
                    parse_companies_payload(
                        data
                    ),
                    "local"
                )

            except Exception as exc:

                local_error = str(exc)

    try:

        with urllib.request.urlopen(
            GITHUB_COMPANIES_URL,
            timeout=12
        ) as response:

            raw = (
                response
                .read()
                .decode("utf-8")
            )

        return (
            parse_companies_payload(
                json.loads(raw)
            ),
            "github"
        )

    except Exception as exc:

        raise RuntimeError(
            "companies.json could not be loaded. "
            f"Local: {local_error}. "
            f"GitHub: {exc}"
        )


# ============================================================
# STRATEGIC ORGANIZATIONS
# ============================================================

STRATEGIC_PROFILES = {

    "golrang": {

        "aliases": [
            "golrang",
            "golrang industrial group",
            "گروه صنعتی گلرنگ",
            "گلرنگ",
        ],

        "label":
            "GOLRANG INDUSTRIAL GROUP",

        "sector":
            "Consumer • Retail • Technology • Investment",

        "accent":
            "gold",

    },


    "mostazafan": {

        "aliases": [
            "mostazafan",
            "foundation of the oppressed",
            "بنیاد مستضعفان",
            "مستضعفان",
        ],

        "label":
            "BONYAD MOSTAZAFAN",

        "sector":
            "Holding • Investment • Industry • Infrastructure",

        "accent":
            "green",

    },


    "copper": {

        "aliases": [
            "national iranian copper",
            "nicico",
            "nicico company",
            "national iranian copper industries",
            "شرکت ملی صنایع مس ایران",
            "صنایع مس ایران",
            "ملی مس",
        ],

        "label":
            "NATIONAL IRANIAN COPPER INDUSTRIES",

        "sector":
            "Mining • Copper • Industry • Investment",

        "accent":
            "blue",

    },

}


def strategic_key(company):

    haystack = " ".join(
        [
            str(
                company.get(
                    "name",
                    ""
                )
            ),

            str(
                company.get(
                    "website",
                    ""
                )
            ),

            str(
                company.get(
                    "sector",
                    ""
                )
            ),

            str(
                company.get(
                    "description",
                    ""
                )
            ),

            str(
                company.get(
                    "country",
                    ""
                )
            ),
        ]
    ).lower()

    for key, profile in (
        STRATEGIC_PROFILES.items()
    ):

        for alias in profile["aliases"]:

            if alias.lower() in haystack:

                return key

    return None


def company_domain(website):

    if not website:
        return ""

    value = re.sub(
        r"^https?://",
        "",
        str(website)
    )

    return value.split("/")[0].strip()


def company_logo_url(company):

    logo = str(
        company.get("logo")
        or company.get("logo_url")
        or ""
    ).strip()

    if logo:
        return logo

    domain = company_domain(
        company.get(
            "website",
            ""
        )
    )

    if domain:

        return (
            "https://www.google.com/"
            "s2/favicons"
            f"?domain={domain}&sz=128"
        )

    return ""


# ============================================================
# GROQ
# ============================================================

def get_groq_client():

    if Groq is None:
        return None

    api_key = None

    try:

        api_key = st.secrets.get(
            "GROQ_API_KEY"
        )

    except Exception:
        pass

    if not api_key:

        api_key = os.getenv(
            "GROQ_API_KEY"
        )

    if not api_key:
        return None

    return Groq(
        api_key=api_key
    )


def ai_language_name():

    return st.session_state.get(
        "language_name",
        "English"
    )


def ask_groq(
    question,
    context,
    mode="qa"
):

    client = get_groq_client()

    if client is None:

        return (
            None,
            t("groq_missing")
        )

    model = os.getenv(
        "GROQ_MODEL",
        "llama-3.3-70b-versatile"
    )

    if mode == "summary":

        system = f"""

You are the SCOCOEX NEXUS
document intelligence assistant.

Summarize ONLY the supplied source material.

Do not invent facts.

Write the answer in
{ai_language_name()}.

Use clear headings,
executive-level language,
and concise bullet points.

If a point is not supported
by the source, do not add it.

"""

        user = (
            "Create a structured executive "
            "summary from this document:\n\n"
            + context
        )

    else:

        system = f"""

You are SCOCOEX NEXUS AI Assistant.

Answer ONLY from the supplied
SCOCOEX Knowledge Base context.

Do not fabricate information.

If the context does not contain
the answer, explicitly say that
the current Knowledge Base does
not provide enough information.

Answer in {ai_language_name()}.

When useful, mention the source
document name.

"""

        user = (
            "Question:\n"
            + question
            + "\n\nKnowledge Base context:\n"
            + context
        )

    try:

        response = (
            client.chat.completions.create(
                model=model,

                messages=[
                    {
                        "role":
                            "system",
                        "content":
                            system.strip(),
                    },

                    {
                        "role":
                            "user",
                        "content":
                            user.strip(),
                    },
                ],

                temperature=0.15,

                max_tokens=1800,
            )
        )

        return (
            response
            .choices[0]
            .message
            .content
            .strip(),
            None
        )

    except Exception as exc:

        return (
            None,
            f"Groq error: {exc}"
        )


# ============================================================
# PDF OUTPUT
# ============================================================

def create_pdf(
    text,
    title="SCOCOEX AI Answer"
):

    if (
        SimpleDocTemplate is None
        or Paragraph is None
    ):

        return None

    buffer = BytesIO()

    styles = getSampleStyleSheet()

    story = [

        Paragraph(
            html.escape(title),
            styles["Title"]
        ),

        Spacer(
            1,
            12
        ),

    ]

    for paragraph in text.split("\n"):

        if paragraph.strip():

            story.append(
                Paragraph(
                    html.escape(
                        paragraph.strip()
                    ),
                    styles["BodyText"]
                )
            )

            story.append(
                Spacer(
                    1,
                    7
                )
            )

    try:

        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4
        )

        doc.build(story)

        return buffer.getvalue()

    except Exception:

        return None


# ============================================================
# LOAD DATA
# ============================================================

documents = load_all_documents()

chunks = create_chunks(
    documents
)

try:

    companies, companies_source = (
        load_companies()
    )

    companies_error = None

except Exception as exc:

    companies = []

    companies_source = None

    companies_error = str(exc)


# ============================================================
# SIDEBAR NAVIGATION
# ============================================================

pages = [

    t("home"),

    t("info"),

    t("assistant"),

    t("companies"),

    t("intel"),

    t("gallery"),

    t("about"),

    t("contact"),

]


page = st.sidebar.radio(

    t("nav"),

    pages,

    key="main_navigation_v4",

    label_visibility="collapsed",

)


st.sidebar.divider()


st.sidebar.caption(
    f"{t('documents')}: "
    f"{len(documents)}"
)


st.sidebar.caption(
    f"{t('international_companies')}: "
    f"{len(companies)}"
)


st.sidebar.caption(
    "SCOCOEX NEXUS • v4.0"
)


# ============================================================
# HOME
# ============================================================

if page == t("home"):

    logo_path = (
        ASSETS_DIR /
        "scocoex_logo-1.png"
    )

    if not logo_path.exists():

        logo_files = sorted(
            ASSETS_DIR.glob(
                "scocoex_logo*.png"
            )
        )

        if logo_files:

            logo_path = logo_files[0]

    if logo_path.exists():

        _, center, _ = st.columns(
            [1, 2, 1]
        )

        with center:

            st.image(
                str(logo_path),
                use_container_width=True
            )

    st.markdown(
        f"""

        <div
            class="hero"
            dir="{direction()}"
        >

            <div class="hero-label">
                {esc(t("hero_label"))}
            </div>

            <h1>
                {t("hero_title")}
            </h1>

            <p>
                {esc(t("hero_text"))}
            </p>

        </div>

        """,

        unsafe_allow_html=True
    )


    cols = st.columns(4)

    metrics = [

        (
            t("documents"),
            len(documents)
        ),

        (
            t("chunks"),
            len(chunks)
        ),

        (
            t("international_companies"),
            len(companies)
        ),

        (
            t("languages"),
            5
        ),

    ]


    for col, (label, value) in zip(
        cols,
        metrics
    ):

        with col:

            st.markdown(
                f"""

                <div
                    class="metric-card"
                    dir="{direction()}"
                >

                    <div class="muted">
                        {esc(label)}
                    </div>

                    <div class="metric-number">
                        {esc(value)}
                    </div>

                </div>

                """,

                unsafe_allow_html=True
            )


    st.markdown(
        f"### {t('knowledge')}"
    )


    cards = [

        (
            "📚",
            t("knowledge"),
            t("knowledge_text")
        ),

        (
            "🤖",
            t("ai"),
            t("ai_text")
        ),

        (
            "🌍",
            t("network"),
            t("network_text")
        ),

    ]


    cols = st.columns(3)


    for col, (
        icon,
        title,
        description
    ) in zip(
        cols,
        cards
    ):

        with col:

            st.markdown(
                f"""

                <div
                    class="feature-card"
                    dir="{direction()}"
                >

                    <div
                        style="font-size:28px;"
                    >
                        {icon}
                    </div>

                    <h3>
                        {esc(title)}
                    </h3>

                    <div
                        class="muted"
                        style="
                            font-size:14px;
                            line-height:1.7;
                        "
                    >
                        {esc(description)}
                    </div>

                </div>

                """,

                unsafe_allow_html=True
            )


# ============================================================
# INFORMATION CENTER
# ============================================================

elif page == t("info"):

    st.markdown(
        f"""
        <div class="section-title">
            {esc(t("information_title"))}
        </div>
        """,

        unsafe_allow_html=True
    )


    st.write(
        t("information_text")
    )


    if not DOCS_DIR.exists():

        st.error(
            t("docs_missing")
        )

    elif not documents:

        st.warning(
            t("no_docs")
        )

    else:

        st.success(
            f"{len(documents)} "
            f"{t('documents').lower()}"
        )


        search = st.text_input(
            t("search_documents"),
            key="document_search_v4"
        )


        filtered_documents = [

            doc

            for doc in documents

            if (
                not search
                or search.lower()
                in doc["name"].lower()
            )

        ]


        for document in (
            filtered_documents
        ):

            with st.expander(
                f"📄 {document['name']}"
            ):

                c1, c2, c3 = (
                    st.columns(3)
                )


                total_words = sum(

                    len(
                        tokenize(
                            section["text"]
                        )
                    )

                    for section
                    in document["sections"]

                )


                full_text = " ".join(

                    section["text"]

                    for section
                    in document["sections"]

                )


                numeric_count = len(

                    re.findall(
                        r"\b\d+(?:[.,]\d+)?\b",
                        full_text
                    )

                )


                c1.metric(
                    t("extracted_sections"),
                    len(
                        document["sections"]
                    )
                )


                c2.metric(
                    t("words"),
                    total_words
                )


                c3.metric(
                    t("numeric_values"),
                    numeric_count
                )


                if document.get(
                    "error"
                ):

                    st.error(
                        document["error"]
                    )

                    continue


                if st.button(
                    t("summary"),
                    key=(
                        "summary_button_v4_"
                        + document["name"]
                    ),
                    use_container_width=True
                ):

                    with st.spinner(
                        t("generating")
                    ):

                        summary, error = (
                            ask_groq(
                                question="",
                                context=
                                    full_text[
                                        :24000
                                    ],
                                mode="summary"
                            )
                        )


                    if error:

                        st.error(error)

                    elif summary:

                        st.markdown(
                            f"""
                            <div
                                class="answer-box"
                                dir="{direction()}"
                            >
                            """,
                            unsafe_allow_html=True
                        )

                        st.markdown(
                            summary
                        )

                        st.markdown(
                            "</div>",
                            unsafe_allow_html=True
                        )


                        pdf_data = create_pdf(

                            summary,

                            "SCOCOEX — "
                            + document["name"]

                        )


                        if pdf_data:

                            st.download_button(

                                t("download_pdf"),

                                data=pdf_data,

                                file_name=(
                                    "scocoex_"
                                    "ai_summary.pdf"
                                ),

                                mime="application/pdf",

                                key=(
                                    "download_summary_v4_"
                                    + document["name"]
                                )

                            )


# ============================================================
# AI ASSISTANT
# ============================================================

elif page == t("assistant"):

    st.markdown(
        f"""
        <div class="section-title">
            {esc(t("ask_title"))}
        </div>
        """,
        unsafe_allow_html=True
    )


    st.write(
        t("ask_caption")
    )


    question = st.text_area(

        t("your_question"),

        placeholder=
            t("your_question")
            + "...",

        height=140,

        key="assistant_question_v4",

    )


    st.caption(
        f"{t('response_language')}: "
        f"{st.session_state.language_name}"
    )


    ask_clicked = st.button(

        t("ask_button"),

        type="primary",

        use_container_width=True,

        key="assistant_ask_button_v4",

    )


    if ask_clicked:

        if not question.strip():

            st.warning(
                t("your_question")
            )

        elif not chunks:

            st.warning(
                t("no_results")
            )

        else:

            with st.spinner(
                t("searching")
            ):

                relevant = (
                    retrieve_chunks(
                        question,
                        chunks,
                        top_k=7
                    )
                )


            if not relevant:

                st.warning(
                    t("no_results")
                )

            else:

                context_parts = []


                for item in relevant:

                    page_info = (

                        f", page "
                        f"{item['page']}"

                        if item.get("page")
                        else ""

                    )


                    context_parts.append(

                        f"SOURCE: "
                        f"{item['document']}"
                        f"{page_info}\n"
                        f"{item['text']}"

                    )


                context = (
                    "\n\n---\n\n"
                    .join(context_parts)
                )


                with st.spinner(
                    t("generating")
                ):

                    answer, error = (
                        ask_groq(
                            question,
                            context,
                            mode="qa"
                        )
                    )


                if error:

                    st.error(error)

                elif answer:

                    st.markdown(
                        f"""
                        <div
                            class="answer-box"
                            dir="{direction()}"
                        >
                        """,
                        unsafe_allow_html=True
                    )

                    st.markdown(
                        answer
                    )

                    st.markdown(
                        "</div>",
                        unsafe_allow_html=True
                    )


                    st.markdown(
                        f"### {t('sources')}"
                    )


                    seen = set()


                    for item in relevant:

                        source_key = (
                            item["document"],
                            item.get("page")
                        )


                        if source_key in seen:
                            continue


                        seen.add(
                            source_key
                        )


                        page_text = (

                            f" — page "
                            f"{item['page']}"

                            if item.get("page")
                            else ""

                        )


                        st.info(
                            "📄 "
                            + item["document"]
                            + page_text
                        )


                    pdf_data = create_pdf(

                        answer,

                        "SCOCOEX NEXUS — "
                        "AI Assistant"

                    )


                    if pdf_data:

                        st.download_button(

                            t("download_pdf"),

                            data=pdf_data,

                            file_name=(
                                "scocoex_ai_answer.pdf"
                            ),

                            mime="application/pdf",

                            key="download_ai_answer_v4",

                        )


# ============================================================
# COMPANIES
# ============================================================

elif page == t("companies"):

    st.markdown(
        f"""
        <div class="section-title">
            {esc(t("companies_title"))}
        </div>
        """,
        unsafe_allow_html=True
    )


    st.markdown(
        """
        <div
            class="muted"
            style="
                font-size:15px;
                margin-bottom:20px;
            "
        >
            SCOCOEX NEXUS strategic organization
            intelligence and global company network.
        </div>
        """,
        unsafe_allow_html=True
    )


    if companies_error:

        st.error(
            companies_error
        )


    if not companies:

        st.warning(
            t("no_companies")
        )

    else:

        st.caption(

            f"{len(companies)} "
            f"{t('company_count')} • "
            f"{t('status')}: "
            f"{t('loaded')} • "
            f"{t('local') if companies_source == 'local' else t('github')}"

        )


        # ----------------------------------------------------
        # STRATEGIC
        # ----------------------------------------------------

        strategic_companies = []

        for company in companies:

            key = strategic_key(
                company
            )

            if key:

                strategic_companies.append(
                    (
                        key,
                        company
                    )
                )


        if strategic_companies:

            st.markdown(
                f"""
                <div
                    class="strategic-header"
                >
                    <span>✦</span>
                    {esc(t("strategic")).upper()}
                </div>
                """,
                unsafe_allow_html=True
            )


            st.caption(
                "Featured strategic organizations "
                "within the SCOCOEX intelligence network."
            )


            strategic_columns = (
                st.columns(3)
            )


            for index, (
                strategic_id,
                company
            ) in enumerate(
                strategic_companies[:3]
            ):

                profile = (
                    STRATEGIC_PROFILES[
                        strategic_id
                    ]
                )


                name = str(
                    company.get(
                        "name"
                    )
                    or profile["label"]
                )


                website = str(
                    company.get(
                        "website"
                    )
                    or ""
                )


                logo = company_logo_url(
                    company
                )


                initial = (
                    name[:1]
                    .upper()
                    if name
                    else "S"
                )


                if logo:

                    logo_html = (

                        '<img '
                        'class="strategic-logo" '
                        f'src="{esc(logo)}" '
                        f'alt="{esc(initial)}">'
                    )

                else:

                    logo_html = (

                        '<div '
                        'class="strategic-logo-fallback">'
                        f'{esc(initial)}'
                        '</div>'
                    )


                link_html = (

                    f'''
                    <a
                        class="strategic-link"
                        href="{esc(website)}"
                        target="_blank"
                        rel="noopener noreferrer"
                    >
                        Explore organization →
                    </a>
                    '''

                    if website

                    else ""
                )


                with strategic_columns[
                    index
                ]:

                    st.markdown(

                        f"""
                        <div
                            class="
                                strategic-card
                                strategic-
                                {profile["accent"]}
                            "
                            dir="{direction()}"
                        >

                            <div
                                class="strategic-top"
                            >

                                {logo_html}

                                <div>

                                    <div
                                        class="
                                            strategic-kicker
                                        "
                                    >
                                        STRATEGIC
                                        PROFILE
                                    </div>

                                    <div
                                        class="
                                            strategic-name
                                        "
                                    >
                                        {esc(name)}
                                    </div>

                                </div>

                            </div>

                            <div
                                class="
                                    strategic-sector
                                "
                            >
                                {esc(
                                    profile["sector"]
                                )}
                            </div>

                            <div
                                class="
                                    strategic-copy
                                "
                            >
                                NEXUS profile for
                                strategic discovery,
                                relationships,
                                opportunities and
                                intelligence.
                            </div>

                            {link_html}

                        </div>
                        """,

                        unsafe_allow_html=True
                    )


        # ----------------------------------------------------
        # INTERNATIONAL NETWORK
        # ----------------------------------------------------

        st.markdown(
            f"""
            <div
                class="directory-header"
            >
                <span>🌍</span>
                {esc(
                    t("international_network")
                ).upper()}
            </div>
            """,

            unsafe_allow_html=True
        )


        all_countries = sorted(
            {
                str(
                    c.get(
                        "country",
                        ""
                    )
                ).strip()

                for c in companies

                if str(
                    c.get(
                        "country",
                        ""
                    )
                ).strip()
            }
        )


        all_sectors = sorted(
            {
                str(
                    c.get(
                        "sector",
                        ""
                    )
                ).strip()

                for c in companies

                if str(
                    c.get(
                        "sector",
                        ""
                    )
                ).strip()
            }
        )


        c1, c2, c3 = st.columns(
            [2.2, 1, 1]
        )


        with c1:

            company_query = st.text_input(

                t("search_company"),

                placeholder=
                    t("search_network"),

                key="company_search_v4",

            )


        with c2:

            country_filter = st.selectbox(

                t("country"),

                [t("all")] + all_countries,

                key="company_country_filter_v4",

            )


        with c3:

            sector_filter = st.selectbox(

                t("sector"),

                [t("all")] + all_sectors,

                key="company_sector_filter_v4",

            )


        filtered = []


        query = (
            company_query
            .strip()
            .lower()
        )


        for company in companies:

            blob = json.dumps(
                company,
                ensure_ascii=False
            ).lower()


            if (
                query
                and query not in blob
            ):

                continue


            if (
                country_filter != t("all")
                and str(
                    company.get(
                        "country",
                        ""
                    )
                )
                != country_filter
            ):

                continue


            if (
                sector_filter != t("all")
                and str(
                    company.get(
                        "sector",
                        ""
                    )
                )
                != sector_filter
            ):

                continue


            filtered.append(
                company
            )


        st.markdown(
            f"""
            <div
                class="directory-count"
            >
                {len(filtered)}
                <span>
                    {esc(
                        t(
                            "organizations_in_network"
                        )
                    )}
                </span>
            </div>
            """,
            unsafe_allow_html=True
        )


        if not filtered:

            st.info(
                "No organizations match "
                "the selected filters."
            )


        company_columns = (
            st.columns(3)
        )


        for index, company in enumerate(
            filtered
        ):

            with company_columns[
                index % 3
            ]:

                name = str(
                    company.get(
                        "name"
                    )
                    or "International Organization"
                )


                website = str(
                    company.get(
                        "website"
                    )
                    or ""
                )


                sector = str(
                    company.get(
                        "sector"
                    )
                    or ""
                )


                country = str(
                    company.get(
                        "country"
                    )
                    or ""
                )


                description = str(
                    company.get(
                        "description"
                    )
                    or ""
                )


                logo = company_logo_url(
                    company
                )


                if logo:

                    logo_html = (

                        '<img '
                        'class="directory-logo" '
                        f'src="{esc(logo)}" '
                        f'alt="{esc(name)}">'
                    )

                else:

                    logo_html = (

                        '<div '
                        'class="directory-initial">'
                        f'{esc(name[:1].upper())}'
                        '</div>'
                    )


                meta = " • ".join(
                    value
                    for value in [
                        country,
                        sector
                    ]
                    if value
                )


                st.markdown(

                    f"""
                    <div
                        class="directory-card"
                        dir="{direction()}"
                    >

                        <div
                            class="directory-top"
                        >

                            {logo_html}

                            <div
                                class="directory-name"
                            >
                                {esc(name)}
                            </div>

                        </div>

                        <div
                            class="directory-meta"
                        >
                            {esc(
                                meta
                                or "International Organization"
                            )}
                        </div>

                        <div
                            class="directory-bottom"
                        >
                            {
                                (
                                    f'<a '
                                    f'class="directory-link" '
                                    f'href="{esc(website)}" '
                                    f'target="_blank" '
                                    f'rel="noopener noreferrer">'
                                    f'{esc(t("visit"))}'
                                    f'</a>'
                                )
                                if website
                                else
                                '<span '
                                'class="directory-link">'
                                'Profile'
                                '</span>'
                            }
                        </div>

                    </div>
                    """,

                    unsafe_allow_html=True
                )


                if description:

                    with st.expander(
                        "View intelligence"
                    ):

                        st.write(
                            description
                        )


# ============================================================
# INTELLIGENCE DASHBOARD
# ============================================================

elif page == t("intel"):

    st.markdown(
        f"""
        <div class="section-title">
            {esc(t("intel_title"))}
        </div>
        """,
        unsafe_allow_html=True
    )


    st.write(
        t("intel_caption")
    )


    if not documents:

        st.warning(
            t("no_docs")
        )

    else:

        rows = []


        for document in documents:

            full_text = " ".join(

                section["text"]

                for section
                in document["sections"]

            )


            rows.append(

                {

                    t("source_file"):
                        document["name"],

                    t("words"):
                        len(
                            tokenize(
                                full_text
                            )
                        ),

                    t("numeric_values"):
                        len(
                            re.findall(
                                r"\b\d+(?:[.,]\d+)?\b",
                                full_text
                            )
                        ),

                    t("extracted_sections"):
                        len(
                            document["sections"]
                        ),

                }

            )


        df = pd.DataFrame(
            rows
        )


        st.dataframe(
            df,
            use_container_width=True,
            hide_index=True
        )


        if not df.empty:

            st.markdown(
                f"### {t('document_size')}"
            )


            st.bar_chart(

                df.set_index(
                    t("source_file")
                )[t("words")]

            )


            st.markdown(
                f"### {t('numeric')}"
            )


            st.bar_chart(

                df.set_index(
                    t("source_file")
                )[t("numeric_values")]

            )


# ============================================================
# GALLERY
# ============================================================

elif page == t("gallery"):

    st.markdown(
        f"""
        <div class="section-title">
            {esc(t("gallery_title"))}
        </div>
        """,
        unsafe_allow_html=True
    )


    st.write(
        t("gallery_text")
    )


    image_extensions = {
        ".png",
        ".jpg",
        ".jpeg",
        ".webp"
    }


    gallery_files = []


    if ASSETS_DIR.exists():

        gallery_files = [

            path

            for path in sorted(
                ASSETS_DIR.rglob("*")
            )

            if (
                path.is_file()
                and
                path.suffix.lower()
                in image_extensions
            )

        ]


    if gallery_files:

        cols = st.columns(3)


        for index, image_path in enumerate(
            gallery_files
        ):

            with cols[
                index % 3
            ]:

                st.image(
                    str(image_path),
                    use_container_width=True
                )

                st.caption(
                    image_path.name
                )

    else:

        st.info(
            t("gallery_text")
        )


# ============================================================
# ABOUT
# ============================================================

elif page == t("about"):

    st.markdown(
        f"""
        <div class="section-title">
            {esc(t("about_title"))}
        </div>
        """,
        unsafe_allow_html=True
    )


    st.markdown(
        f"""
        <div
            class="feature-card"
            dir="{direction()}"
        >

            <h3>
                SCOCOEX NEXUS
            </h3>

            <p
                style="
                    color:#B8C0D4;
                    line-height:1.8;
                "
            >
                {esc(t("hero_text"))}
            </p>

        </div>
        """,
        unsafe_allow_html=True
    )


# ============================================================
# CONTACT
# ============================================================

elif page == t("contact"):

    st.markdown(
        f"""
        <div class="section-title">
            {esc(t("contact_title"))}
        </div>
        """,
        unsafe_allow_html=True
    )


    with st.form(
        "scocex_contact_form_v4"
    ):

        name = st.text_input(
            t("name"),
            key="contact_name_v4"
        )


        company = st.text_input(
            t("company_org"),
            key="contact_company_v4"
        )


        email = st.text_input(
            t("email"),
            key="contact_email_v4"
        )


        interest = st.selectbox(

            t("interest"),

            [
                "Investment",
                "Partnership",
                "International Company",
                "Government",
                "Technology",
                "Media",
                "Other",
            ],

            key="contact_interest_v4"

        )


        message = st.text_area(

            t("message"),

            height=160,

            key="contact_message_v4"

        )


        submitted = (
            st.form_submit_button(
                t("send"),
                type="primary",
                use_container_width=True
            )
        )


        if submitted:

            st.success(
                t("received")
            )


# ============================================================
# FOOTER
# ============================================================

st.markdown(
    """
    <div
        style="
            margin-top:45px;
            padding-top:18px;
            border-top:
                1px solid
                rgba(255,255,255,.07);
            color:#687289;
            font-size:11px;
            letter-spacing:.08em;
            text-align:center;
        "
    >
        SCOCOEX NEXUS
        • AI KNOWLEDGE PLATFORM
        • v4.0
    </div>
    """,
    unsafe_allow_html=True
)
